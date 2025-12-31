"""
Generate publication-ready DOCX with embedded tables and figures
Author: Dr. Siddalingaiah H S
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE_DIR = Path(__file__).parent

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_formatted_run(para, text):
    """Add text with superscript handling"""
    parts = re.split(r'(\^\d+(?:,\d+)*\^)', text)
    for part in parts:
        if part.startswith('^') and part.endswith('^'):
            run = para.add_run(part[1:-1])
            run.font.superscript = True
        else:
            para.add_run(part)

def create_manuscript():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('', level=0)
    title_run = title.add_run('An Integrated Multi-omics and Chemoinformatics Pipeline Identifies IDO1, PDL1, and JAK2 as Host-Directed Therapy Candidates for Leprosy')
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Running title
    rt = doc.add_paragraph()
    rt.add_run('Running Title: ').bold = True
    rt.add_run('Host-Directed Therapy Targets for Leprosy')
    
    doc.add_paragraph()
    
    # Authors
    authors = doc.add_paragraph()
    authors.add_run('Authors: ').bold = True
    run = authors.add_run('Siddalingaiah H S')
    sup = authors.add_run('1*')
    sup.font.superscript = True
    
    doc.add_paragraph()
    
    # Affiliations
    aff = doc.add_paragraph()
    aff.add_run('Affiliations: ').bold = True
    sup = aff.add_run('1')
    sup.font.superscript = True
    aff.add_run('Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur – 572106, Karnataka, India')
    
    doc.add_paragraph()
    
    # Corresponding author
    corr = doc.add_paragraph()
    corr.add_run('*Corresponding Author: ').bold = True
    corr.add_run('Dr. Siddalingaiah H S, Professor, Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur – 572106, Karnataka, India. Email: hssling@yahoo.com; Phone: +91-8941087719; ORCID: 0000-0002-4771-8285')
    
    doc.add_paragraph()
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run('Word Count: ').bold = True
    meta.add_run('3,850 words | ')
    meta.add_run('Tables: ').bold = True
    meta.add_run('3 | ')
    meta.add_run('Figures: ').bold = True
    meta.add_run('4 | ')
    meta.add_run('References: ').bold = True
    meta.add_run('35')
    
    doc.add_page_break()
    
    # ===== ABSTRACT =====
    doc.add_heading('ABSTRACT', level=1)
    
    abstract_sections = [
        ('Background:', 'Leprosy (Hansen\'s disease), caused by Mycobacterium leprae, remains endemic in tropical countries with approximately 200,000 new cases annually. Despite effective multidrug therapy (MDT), leprosy reactions cause significant nerve damage and disability. Host-directed therapies (HDTs) represent an emerging strategy to modulate immune responses.'),
        ('Objectives:', 'To systematically identify druggable host targets and repurposable drug candidates for leprosy using an integrated multi-omics and chemoinformatics pipeline.'),
        ('Methods:', 'A 50-gene host signature was compiled from published leprosy transcriptomic studies (GEO datasets GSE16844, GSE125943, GSE129033, GSE74481). An automated Python pipeline integrated MyGene.info, Open Targets Platform, and ChEMBL database for target prioritization and compound mining.'),
        ('Results:', 'Fifty host targets were prioritized. Top candidates were VEGFA (score 0.525), PDL1/CD274 (0.449), IDO1 (0.386), CD38 (0.291), and IL10 (0.273). A total of 629 bioactive compounds (pChEMBL ≥6.0) were identified, including FDA-approved JAK inhibitors. Literature validation confirmed IDO1\'s role in immunosuppression, PDL1-mediated T-cell exhaustion, and Tofacitinib efficacy for chronic ENL.'),
        ('Conclusions:', 'IDO1, PDL1, and JAK2 emerged as high-priority HDT targets with strong literature validation. These findings support developing adjunctive immunotherapies beyond conventional MDT.')
    ]
    
    for label, text in abstract_sections:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(' ' + text)
    
    # Keywords
    kw = doc.add_paragraph()
    kw.add_run('Keywords: ').bold = True
    kw.add_run('leprosy; host-directed therapy; IDO1; PD-L1; immune checkpoint; JAK2; drug repurposing; Mycobacterium leprae; tuberculosis')
    
    doc.add_page_break()
    
    # ===== INTRODUCTION =====
    doc.add_heading('1. INTRODUCTION', level=1)
    
    intro_paras = [
        'Leprosy (Hansen\'s disease), a chronic granulomatous infection caused by the obligate intracellular pathogen Mycobacterium leprae, remains a significant public health concern despite being declared "eliminated as a public health problem" by the World Health Organization (WHO) in 2000.^1,2^ Approximately 200,000 new cases are registered annually, with India, Brazil, and Indonesia collectively accounting for over 80% of the global burden.^3^ The disease disproportionately affects marginalized populations in tropical and subtropical regions, perpetuating cycles of poverty and social stigma that persist long after microbiological cure.^4^',
        
        'The clinical spectrum of leprosy reflects the host\'s immune response to M. leprae, ranging from the paucibacillary tuberculoid (TT) form with robust cell-mediated immunity and few bacilli to the multibacillary lepromatous (LL) form characterized by anergic responses and high bacterial loads.^5,6^ This immunological dichotomy, first described by Ridley and Jopling in 1966, remains the foundation for understanding leprosy pathogenesis. Between these poles lie borderline forms (BT, BB, BL) with intermediate immunological profiles, reflecting the dynamic interplay between host immunity and pathogen persistence.',
        
        'A critical challenge in leprosy management is the occurrence of leprosy reactions—immunologically mediated inflammatory episodes affecting up to 50% of patients during or after treatment.^7^ Type 1 reactions (reversal reactions, RR) result from delayed-type hypersensitivity responses as cell-mediated immunity improves with treatment, causing acute inflammation in skin lesions and nerves. Type 2 reactions (erythema nodosum leprosum, ENL) represent immune complex-mediated phenomena with systemic manifestations including fever, painful subcutaneous nodules, and neuritis.^8,9^ Both reaction types cause acute nerve damage leading to irreversible disability even after successful antimicrobial treatment, representing the primary cause of leprosy-associated morbidity in the post-MDT era.',
        
        'Standard multidrug therapy (MDT) combining dapsone, rifampicin, and clofazimine has revolutionized leprosy treatment, reducing prevalence by over 90% since 1985.^10^ However, MDT does not prevent reactions, and current immunosuppressive treatments—corticosteroids and thalidomide—have significant adverse effects including metabolic complications, opportunistic infections, and teratogenicity.^11,12^ Moreover, 10-20% of ENL cases become chronic and steroid-refractory, necessitating prolonged immunosuppression with attendant complications. This therapeutic gap underscores the urgent need for novel strategies that can modulate host immune responses more precisely without broad immunosuppression.',
        
        'Host-directed therapies (HDTs) represent an emerging paradigm in infectious disease treatment, targeting host cellular pathways rather than pathogen-specific mechanisms.^13,14^ Originally pioneered for tuberculosis, HDT approaches have demonstrated success in modulating granuloma dynamics, enhancing autophagy, and optimizing inflammatory responses without compromising antimicrobial immunity.^15,16^ The rationale for HDT in mycobacterial infections is compelling: these pathogens have evolved over millennia to manipulate host immunity, creating niches for intracellular survival and persistence. By targeting the host pathways exploited by pathogens, HDTs offer potential advantages including reduced selection for drug resistance, applicability across drug-resistant strains, and synergy with conventional antimicrobials.',
        
        'Mycobacterium leprae, sharing substantial genomic homology with M. tuberculosis despite massive gene decay, employs similar immune evasion mechanisms including phagosome maturation arrest, inhibition of autophagy, and modulation of pattern recognition receptor signaling.^17^ This suggests that HDT strategies validated for tuberculosis may be transferable to leprosy. However, leprosy presents unique challenges: M. leprae cannot be cultured in vitro, limiting drug screening; the armadillo is the only laboratory animal model for disseminated infection; and the extremely slow replication rate (doubling time ~14 days) necessitates prolonged preclinical studies.',
        
        'Computational approaches integrating multi-omics data with chemical databases offer a powerful method to overcome these limitations and accelerate target identification. We previously developed and validated an automated informatics pipeline for identifying HDT targets in tuberculosis, demonstrating its ability to prioritize druggable candidates and identify repurposable compounds with clinical evidence.^21^ In this study, we extended this pipeline to leprosy, integrating published transcriptomic signatures from four independent GEO datasets with druggability assessments from Open Targets Platform and compound bioactivity data from ChEMBL database to systematically identify host targets amenable to pharmacological intervention.'
    ]
    
    for text in intro_paras:
        p = doc.add_paragraph()
        add_formatted_run(p, text)
    
    # ===== METHODS =====
    doc.add_heading('2. MATERIALS AND METHODS', level=1)
    
    doc.add_heading('2.1 Study Design', level=2)
    p = doc.add_paragraph()
    add_formatted_run(p, 'This computational, in silico study employed a systems biology approach to integrate publicly available leprosy transcriptomic data with chemical-genomic databases.^22^ The analysis was conducted in accordance with FAIR principles.')
    
    doc.add_heading('2.2 Host Gene Signature Curation', level=2)
    p = doc.add_paragraph()
    add_formatted_run(p, 'A 50-gene leprosy host signature was compiled from published transcriptomic studies and Gene Expression Omnibus (GEO) datasets.^23^ The following datasets were utilized: GSE16844 (skin lesion transcriptomes comparing leprosy patients with controls),^24^ GSE125943 (whole blood signatures during leprosy reactions),^25^ GSE129033 (peripheral blood mononuclear cell profiles across the leprosy spectrum),^26^ and GSE74481 (skin transcriptomes during immune reactions).^27^ This multi-dataset approach enabled identification of genes consistently dysregulated across diverse sample types and clinical presentations.')
    
    p = doc.add_paragraph()
    add_formatted_run(p, 'Genes were selected based on three criteria: (1) consistent differential expression across studies (|log₂ fold change| ≥1.0), (2) statistical significance (adjusted p-value <0.05), and (3) biological relevance to leprosy immunopathology based on literature review. The final signature encompassed genes involved in immune checkpoint regulation (PDL1, CTLA4, PD1), cytokine signaling (JAK2, STAT1, STAT3, IL6, IL10, TNF), tryptophan metabolism (IDO1), pattern recognition (TLR2, TLR10, NOD2), and tissue remodeling (MMP9, VEGFA). This curated approach balanced statistical rigor with biological plausibility, prioritizing genes with established roles in mycobacterial immunity.')
    
    doc.add_heading('2.3 Computational Pipeline', level=2)
    p = doc.add_paragraph()
    add_formatted_run(p, 'The automated pipeline (Python 3.12) integrated: MyGene.info for gene-protein mapping,^28^ Open Targets Platform for druggability assessment,^29^ and ChEMBL database (v33) for compound bioactivity mining.^30^')
    
    doc.add_heading('2.4 Target Prioritization Algorithm', level=2)
    p = doc.add_paragraph()
    p.add_run('Composite Score = 0.35 × Omics_Strength + 0.25 × OpenTargets_Evidence + 0.20 × Druggability_Proxy + 0.10 × Pathway_Centrality + 0.10 × Replication').italic = True
    
    doc.add_page_break()
    
    # ===== RESULTS =====
    doc.add_heading('3. RESULTS', level=1)
    
    doc.add_heading('3.1 Host Target Prioritization', level=2)
    p = doc.add_paragraph()
    add_formatted_run(p, 'The pipeline successfully processed the 50-gene leprosy signature. Scores ranged from 0.013 to 0.525, with median 0.113. The top 10 targets are presented in Table 1.')
    
    # ----- TABLE 1 -----
    doc.add_paragraph()
    table1_title = doc.add_paragraph()
    table1_title.add_run('Table 1: Top 10 Host-Directed Therapy Target Candidates for Leprosy').bold = True
    
    table1 = doc.add_table(rows=11, cols=5)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    headers = ['Rank', 'Gene', 'Protein', 'Score', 'Leprosy Relevance']
    for i, h in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    # Data
    data = [
        ('1', 'VEGFA', 'Vascular Endothelial Growth Factor A', '0.525', 'Granuloma angiogenesis'),
        ('2', 'PDL1', 'Programmed Death Ligand 1', '0.449', 'T-cell exhaustion'),
        ('3', 'IDO1', 'Indoleamine 2,3-Dioxygenase 1', '0.386', 'Tryptophan catabolism'),
        ('4', 'CD38', 'CD38 Molecule', '0.291', 'B-cell marker'),
        ('5', 'IL10', 'Interleukin 10', '0.273', 'Anti-inflammatory'),
        ('6', 'VDR', 'Vitamin D Receptor', '0.271', 'Antimicrobial peptides'),
        ('7', 'BLK', 'B Lymphoid Kinase', '0.270', 'B-cell signaling'),
        ('8', 'JAK2', 'Janus Kinase 2', '0.263', 'Cytokine signaling'),
        ('9', 'CXCL10', 'C-X-C Chemokine Ligand 10', '0.263', 'T-cell recruitment'),
        ('10', 'CTLA4', 'Cytotoxic T-Lymphocyte Antigen 4', '0.259', 'Immune checkpoint'),
    ]
    
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            table1.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # ----- FIGURE 1 -----
    fig1_para = doc.add_paragraph()
    fig1_para.add_run('Figure 1: Top 20 Host-Directed Therapy Targets for Leprosy').bold = True
    doc.add_picture(str(BASE_DIR / 'outputs' / 'figures' / 'figure1_target_prioritization.png'), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Compound Discovery', level=2)
    p = doc.add_paragraph()
    add_formatted_run(p, 'For targets overlapping between the leprosy signature and our previously validated tuberculosis pipeline, 629 bioactive compounds with pChEMBL ≥6.0 (corresponding to IC₅₀ ≤1 μM) were identified from ChEMBL database. The distribution of compound potency is shown in Figure 2B, with 89 compounds (14.1%) demonstrating sub-nanomolar activity (pChEMBL ≥9.0, IC₅₀ ≤1 nM). The target-by-potency analysis (Figure 2C) identified Matrix Metalloproteinase 9 (MMP9), JAK2, and PPARG as targets with the highest-potency available compounds.')
    
    p = doc.add_paragraph()
    add_formatted_run(p, 'Notably, 127 compounds (20.2%) have reached clinical development, with 45 FDA-approved drugs identified as potential repurposing candidates. The clinical phase distribution included: Phase 4 (approved, n=45), Phase 3 (n=23), Phase 2 (n=31), Phase 1 (n=18), and preclinical (n=10). This high proportion of clinically advanced compounds reflects the druggability of prioritized targets and provides a robust pipeline for translational development. Key drug candidates with published clinical evidence are shown in Table 2.')
    
    # ----- TABLE 2 -----
    doc.add_paragraph()
    table2_title = doc.add_paragraph()
    table2_title.add_run('Table 2: Clinically Advanced Drug Candidates for Leprosy HDT').bold = True
    
    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ['Drug', 'Target', 'pChEMBL', 'Phase', 'Evidence']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    data2 = [
        ('Tofacitinib', 'JAK1/2/3', '7.41', '4 (Approved)', 'Published for ENL'),
        ('Ruxolitinib', 'JAK1/2', '7.51', '4 (Approved)', 'Immunomodulation'),
        ('Baricitinib', 'JAK1/2', '7.80', '4 (Approved)', 'Anti-inflammatory'),
        ('Upadacitinib', 'JAK1/2', '8.52', '4 (Approved)', 'High potency'),
        ('Ilomastat', 'MMP9', '10.30', '2', 'Tissue protection'),
    ]
    
    for i, row_data in enumerate(data2):
        for j, val in enumerate(row_data):
            table2.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # ----- FIGURE 2 -----
    fig2_para = doc.add_paragraph()
    fig2_para.add_run('Figure 2: Distribution of Compound Bioactivity').bold = True
    doc.add_picture(str(BASE_DIR / 'outputs' / 'figures' / 'figure2_compound_distribution.png'), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    doc.add_heading('3.3 Literature Validation', level=2)
    
    p = doc.add_paragraph()
    p.add_run('IDO1 (Rank 3): ').bold = True
    add_formatted_run(p, 'IDO1 is robustly validated in leprosy literature. De Mattos Barbosa et al. demonstrated that M. leprae induces IDO1 expression in monocytes via IL-10-dependent mechanisms.^18^ IDO1 expression is significantly elevated in lepromatous lesions.^33^')
    
    p = doc.add_paragraph()
    p.add_run('PDL1/CD274 (Rank 2): ').bold = True
    add_formatted_run(p, 'PD-1/PD-L1 signaling is critically implicated in leprosy immunopathology. Palermo et al. demonstrated elevated PD-1 expression on T cells in lepromatous patients.^19^ In vitro blockade restores IFN-γ production.^34^')
    
    p = doc.add_paragraph()
    p.add_run('JAK2 (Rank 8): ').bold = True
    add_formatted_run(p, 'Thangaraju et al. reported successful use of Tofacitinib for chronic, steroid-refractory ENL.^32^')
    
    # ----- TABLE 3 -----
    doc.add_paragraph()
    table3_title = doc.add_paragraph()
    table3_title.add_run('Table 3: Literature Validation Summary').bold = True
    
    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['Target', 'PubMed Studies', 'Key Finding', 'Therapeutic Implication']
    for i, h in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')
    
    data3 = [
        ('IDO1', '12', 'Upregulated in LL lesions', 'IDO inhibitors may restore immunity'),
        ('PDL1', '8', 'T-cell exhaustion marker', 'Checkpoint blockade potential'),
        ('JAK2', '3', 'Cytokine signaling node', 'Tofacitinib for ENL'),
        ('VDR', '15', 'Cathelicidin induction', 'Vitamin D supplementation'),
        ('MMP9', '7', 'Tissue damage marker', 'MMP inhibitors for nerve protection'),
    ]
    
    for i, row_data in enumerate(data3):
        for j, val in enumerate(row_data):
            table3.rows[i+1].cells[j].text = val
    
    doc.add_paragraph()
    
    # ----- FIGURE 3 -----
    fig3_para = doc.add_paragraph()
    fig3_para.add_run('Figure 3: Top 15 Targets by Maximum Compound Potency').bold = True
    doc.add_picture(str(BASE_DIR / 'outputs' / 'figures' / 'figure3_target_potency.png'), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ----- FIGURE 4 -----
    fig4_para = doc.add_paragraph()
    fig4_para.add_run('Figure 4: Key Leprosy HDT Targets by Pathway').bold = True
    doc.add_picture(str(BASE_DIR / 'outputs' / 'figures' / 'figure4_pathway_heatmap.png'), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ===== DISCUSSION =====
    doc.add_heading('4. DISCUSSION', level=1)
    
    discussion_paras = [
        'This study demonstrates the utility of an integrated multi-omics and chemoinformatics pipeline for systematically identifying host-directed therapy candidates in leprosy. By combining transcriptomic signatures with druggability assessments and compound bioactivity data, we prioritized 50 host targets and identified 629 bioactive compounds, including multiple FDA-approved drugs suitable for repurposing. The pipeline successfully integrated data from four independent GEO datasets, encompassing diverse sample types (skin lesions, whole blood, PBMCs) and clinical presentations across the leprosy spectrum. This comprehensive approach enabled identification of targets relevant to both disease pathogenesis and immunological complications.',
        
        'The prioritization algorithm weighted multiple evidence streams, with omics strength (35%) and Open Targets druggability evidence (25%) contributing most heavily to composite scores. This weighting reflects the dual imperatives of biological relevance and therapeutic tractability. Targets with high composite scores demonstrated convergent evidence from transcriptomic dysregulation, established druggability, and literature validation, providing confidence in their potential as HDT candidates.',
        
        'IDO1 emerged as the third-ranked target with compelling literature validation spanning mechanistic, expression, and functional studies. The IDO1-kynurenine pathway represents a critical mechanism by which M. leprae establishes immune tolerance in the host.^18^ In lepromatous leprosy, elevated IDO1 expression in lesional macrophages and Schwann cells depletes local tryptophan, starving pathogen-reactive T cells while generating immunosuppressive kynurenine metabolites that promote regulatory T-cell differentiation. This mechanism parallels findings in tuberculosis granulomas, where IDO1 activity correlates with disease progression and treatment failure.^21^',
        
        'IDO1 inhibitors, including epacadostat (INCB024360) and indoximod (NLG-8189), have been developed for oncology applications based on the rationale that tumors exploit IDO1 to evade immune surveillance.^15^ While clinical trials in cancer have yielded mixed results, the infectious disease context may differ fundamentally. In cancer, IDO1 inhibition aims to unleash pre-existing anti-tumor immunity; in leprosy, it could reverse pathogen-induced tolerogenic programming. However, the dual role of IDO1 in both antimicrobial defense (through tryptophan starvation of intracellular pathogens) and tissue protection (limiting immunopathology) necessitates careful dose-finding and patient selection in preclinical studies.',
        
        'The identification of PDL1 (CD274) and CTLA4 among top targets aligns with emerging evidence that M. leprae exploits immune checkpoint pathways to evade host immunity.^19^ Palermo et al. demonstrated that PD-1 expression on CD4+ and CD8+ T cells is significantly elevated in lepromatous patients compared to tuberculoid patients or healthy controls, correlating with disease severity and bacterial index. In vitro blockade of PD-1/PD-L1 interaction restores IFN-γ production and T-cell proliferation, suggesting that checkpoint inhibitors could reverse T-cell exhaustion in leprosy.^34^',
        
        'Checkpoint inhibitors (pembrolizumab, nivolumab, ipilimumab) have transformed oncology, achieving durable responses in previously untreatable malignancies. However, their application in infectious diseases remains investigational, with limited data from chronic viral infections (HIV, HCV) and emerging interest in tuberculosis.^16^ In leprosy, potential benefits include restoring antimicrobial immunity in anergic lepromatous patients who fail to mount effective cell-mediated responses. However, excessive immune activation could theoretically precipitate or exacerbate leprosy reactions—the very complications HDT aims to prevent. This paradox necessitates careful patient stratification, potentially restricting checkpoint blockade to paucibacillary patients or combining it with anti-inflammatory agents.',
        
        'Our identification of JAK2 is particularly noteworthy given published clinical evidence supporting JAK inhibition for leprosy reactions.^32^ Thangaraju et al. reported successful use of Tofacitinib (5 mg twice daily) for chronic, steroid-refractory ENL in a case series, representing the first published use of JAK inhibitors in leprosy. The compound\'s anti-inflammatory effects dampen excessive TNF-α, IL-6, and IL-1β production characteristic of ENL, while avoiding the metabolic and infectious complications of chronic corticosteroid use. Our compound analysis identified multiple FDA-approved JAK inhibitors with sub-micromolar potency, including newer selective agents (Upadacitinib, Filgotinib) that may offer improved safety profiles through reduced off-target effects.',
        
        'The JAK-STAT pathway mediates signaling for numerous cytokines implicated in leprosy immunopathology, including IFN-γ, IL-6, IL-10, IL-12, and IL-23. Selective JAK inhibition could theoretically modulate specific cytokine axes while preserving others. For example, JAK1/JAK2 inhibition would dampen IFN-γ and IL-6 signaling (potentially beneficial in ENL) while sparing JAK3-dependent IL-2 signaling (important for T-cell homeostasis). However, broad immunosuppression raises concerns about reactivation of latent infections, particularly tuberculosis in endemic regions where TB-leprosy co-infection occurs.',
        
        'The substantial overlap between leprosy and tuberculosis HDT targets underscores the shared mycobacterial host-pathogen interface. Both pathogens induce IDO1, exploit PD-1/PD-L1 signaling, drive MMP-mediated tissue damage, and modulate vitamin D receptor signaling.^16,35^ This convergence has important translational implications. First, HDT candidates validated for tuberculosis (where larger patient populations and research infrastructure enable clinical trials) may accelerate leprosy drug development through repurposing. Second, combined anti-mycobacterial HDT strategies could benefit patients with mixed infections in endemic regions. Third, shared mechanistic insights may inform development of pan-mycobacterial host-targeted interventions applicable across the Mycobacterium genus.',
        
        'Beyond the top-ranked targets, several additional candidates warrant discussion. VEGFA (rank 1, score 0.525) regulates angiogenesis in granulomas, with potential implications for drug penetration and immune cell trafficking. VDR (rank 6, score 0.271) mediates vitamin D-induced expression of cathelicidin (LL-37), an antimicrobial peptide with direct activity against M. leprae. Observational studies have associated vitamin D deficiency with leprosy susceptibility, and supplementation trials have shown modest benefits.^16^ MMP9 (rank 15, score 0.105) contributes to nerve damage through extracellular matrix degradation, making MMP inhibitors attractive for neuroprotection. However, MMP inhibitors have historically failed in clinical development due to musculoskeletal toxicity, necessitating novel delivery strategies (e.g., topical formulations for cutaneous lesions).'
    ]
    
    for text in discussion_paras:
        p = doc.add_paragraph()
        add_formatted_run(p, text)
    
    # ===== NEW SECTION 4.1: IMMUNOLOGICAL CONTEXT (Peer Review Addition) =====
    doc.add_heading('4.1 Immunological Context and HDT Timing', level=2)
    
    immunology_paras = [
        'Leprosy reactions occur in distinct immunological contexts that have critical implications for HDT selection and timing. Type 1 reactions (reversal reactions) involve upregulation of Th1 responses with increased IFN-γ, TNF-α, and IL-12 production as cell-mediated immunity strengthens during treatment. In this context, checkpoint inhibitors such as anti-PD-1/PD-L1 antibodies could potentially exacerbate inflammation and should be reserved for post-reaction periods or for lepromatous patients with persistent anergy. Conversely, Type 2 reactions (ENL) are characterized by immune complex deposition, complement activation, and excessive proinflammatory cytokine production (particularly TNF-α and IL-6). JAK inhibitors are mechanistically appropriate for ENL as they directly dampen cytokine signaling cascades.',
        
        'The timing of HDT administration requires careful consideration. We propose a stratified approach: (1) Prevention—during MDT for high-risk lepromatous patients with high bacterial index who are at greatest risk for reactions; (2) Treatment—JAK inhibitors for active steroid-refractory ENL as demonstrated by Thangaraju et al.^32^; (3) Maintenance—checkpoint modulation post-MDT completion for lepromatous patients to prevent relapse and enhance bacterial clearance. IDO1 inhibition may be contraindicated during active reactions as it could trigger immune reconstitution inflammatory syndrome (IRIS), similar to that observed in HIV/TB co-infection.^16^ Patient stratification by leprosy type, reaction history, and bacterial load is essential for safe HDT implementation.'
    ]
    
    for text in immunology_paras:
        p = doc.add_paragraph()
        add_formatted_run(p, text)
    
    # ===== NEW SECTION 4.2: SAFETY CONSIDERATIONS (Peer Review Addition) =====
    doc.add_heading('4.2 Safety Considerations', level=2)
    
    safety_paras = [
        'Checkpoint Inhibitors: PD-1/PD-L1 blockade carries well-documented autoimmune risks, with immune-related adverse events (colitis, pneumonitis, hepatitis, thyroiditis) occurring in 10-20% of cancer patients receiving these agents. In the leprosy context, an additional concern is the potential to precipitate Type 1 reactions by enhancing Th1 responses. Mitigation strategies include excluding patients with autoimmune history, restricting use to lepromatous patients post-MDT, initiating at low doses with gradual escalation, and combining with corticosteroid prophylaxis. Close monitoring for both reaction symptoms and autoimmune manifestations is essential.',
        
        'IDO1 Inhibitors: Epacadostat, the most clinically advanced IDO1 inhibitor, failed to improve outcomes in the Phase III ECHO-301 melanoma trial, raising questions about the therapeutic tractability of this target. However, the chronic infection context differs from acute cancer immunotherapy. A critical concern is IDO1\'s dual role: antimicrobial (through tryptophan starvation of intracellular pathogens) versus immunosuppressive (through kynurenine-mediated Treg induction). In leprosy, the balance appears tilted toward immunosuppression in lepromatous lesions, suggesting potential benefit from inhibition. Recommendation: IDO1 inhibition should be considered only post-MDT when bacterial load is controlled, with extensive preclinical testing in the armadillo model before human trials.',
        
        'JAK Inhibitors: The most immediate safety concern is the well-documented risk of tuberculosis reactivation with Tofacitinib and Baricitinib, with TB occurring at 3-4 times the background rate in clinical trials. In TB-endemic areas (India, Brazil, Indonesia) where leprosy is also prevalent, this is a critical consideration. Mitigation requires TB screening (IGRA, chest X-ray) before initiation and regular symptom surveillance during treatment. Additionally, rifampicin (a standard MDT component) is a potent CYP3A4 inducer that reduces JAK inhibitor exposure by 50-70%, potentially compromising efficacy. Solutions include sequential therapy (HDT after MDT completion) or dose adjustment with pharmacokinetic studies to determine optimal dosing in the presence of rifampicin.'
    ]
    
    for text in safety_paras:
        p = doc.add_paragraph()
        add_formatted_run(p, text)
    
    # ===== NEW SECTION 4.3: VEGFA DISCUSSION (Peer Review Addition) =====
    doc.add_heading('4.3 Computational Scoring versus Biological Plausibility', level=2)
    
    p = doc.add_paragraph()
    add_formatted_run(p, 'VEGFA ranked first in our prioritization (score 0.525), primarily driven by its high druggability score reflecting 1,193 known drugs targeting this pathway (predominantly oncology applications). However, biological plausibility for leprosy HDT is limited: only 4 PubMed publications link VEGFA to leprosy, with no functional studies in M. leprae models. Angiogenesis in granulomas is complex and context-dependent—both pro-inflammatory (enabling immune cell infiltration) and potentially beneficial for drug penetration. This discrepancy between computational rank and literature validation highlights a critical limitation: computational scores based on druggability do not equate to biological priority for a specific disease. We explicitly recommend that future studies prioritize IDO1, PDL1, and JAK2 based on their stronger literature validation over VEGFA despite lower computational scores. Future pipeline iterations should incorporate disease-specific biological plausibility weighting to better align computational output with translational potential.')
    
    p = doc.add_paragraph()
    add_formatted_run(p, 'Similarly, BLK (B-lymphoid kinase) ranked 7th (score 0.270) despite having zero leprosy-specific publications. This represents a computational artifact where high general druggability inflates scores for targets with no disease-specific validation. Researchers using similar computational approaches should interpret high-ranking targets with caution and always perform systematic literature validation before advancing candidates to experimental studies.')
    
    # ===== NEW SECTION 4.4: LIMITATIONS (Peer Review Addition) =====
    doc.add_heading('4.4 Limitations', level=2)
    
    p = doc.add_paragraph()
    add_formatted_run(p, 'This study has several important limitations. First, computational predictions, while systematic and reproducible, require experimental validation in appropriate disease models. Second, M. leprae cannot be cultured in vitro, severely limiting drug screening capabilities and necessitating reliance on the armadillo model, which is expensive, slow (requiring 12-18 months for disease development), and ethically complex. Third, the composite scoring algorithm, while incorporating multiple evidence streams, uses weights that, despite sensitivity analysis, remain somewhat arbitrary. Fourth, literature validation depends on published studies and may miss recent unpublished findings or negative results. Fifth, pathway centrality was calculated from general protein-protein interaction networks (STRING v11.5) rather than leprosy-specific interactomes, which may not accurately reflect disease-specific pathway importance. Sixth, the identification of targets like BLK with high computational scores but zero disease-specific publications demonstrates that the pipeline can generate artifacts requiring manual curation. Finally, the overlap with tuberculosis targets, while potentially advantageous for repurposing, limits identification of leprosy-specific host factors that may be uniquely important for this disease.')
    
    # ===== CONCLUSIONS =====
    doc.add_heading('5. CONCLUSIONS', level=1)
    p = doc.add_paragraph()
    add_formatted_run(p, 'This study presents a systematic, reproducible computational approach for identifying host-directed therapy candidates in leprosy. We prioritized IDO1, PDL1, and JAK2 as high-priority druggable targets with strong literature validation. The identification of FDA-approved JAK inhibitors as repurposable candidates provides a clear path for translational development. These findings support developing adjunctive immunotherapies to complement conventional MDT and prevent leprosy-associated disability.')
    
    # ===== ACKNOWLEDGEMENTS =====
    doc.add_heading('ACKNOWLEDGEMENTS', level=1)
    doc.add_paragraph('The author acknowledges the ChEMBL team at EMBL-EBI for open-access bioactivity databases, the Open Targets consortium for druggability assessments, and the GEO database for hosting publicly accessible transcriptomic data.')
    
    doc.add_page_break()
    
    # ===== REFERENCES =====
    doc.add_heading('REFERENCES', level=1)
    
    references = [
        'World Health Organization. Leprosy elimination: an operational manual. Geneva: WHO; 2016.',
        'Lockwood DN, Suneetha S. Leprosy: too complex a disease for a simple elimination paradigm. Bull World Health Organ 2005; 83: 230-235.',
        'World Health Organization. Global leprosy update, 2022: moving towards a leprosy-free world. Wkly Epidemiol Rec 2023; 98: 409-430.',
        'Sermrittirong S, Van Brakel WH. How to reduce stigma in leprosy. Lepr Rev 2014; 85: 149-157.',
        'Ridley DS, Jopling WH. Classification of leprosy according to immunity. Int J Lepr 1966; 34: 255-273.',
        'Scollard DM, et al. The continuing challenges of leprosy. Clin Microbiol Rev 2006; 19: 338-381.',
        'Walker SL, et al. Development and validation of a severity scale for leprosy type 1 reactions. PLoS Negl Trop Dis 2008; 2: e351.',
        'Walker SL, Lockwood DN. Leprosy type 1 (reversal) reactions. Lepr Rev 2008; 79: 372-386.',
        'Kahawita IP, et al. Leprosy type 1 reactions and ENL. An Bras Dermatol 2008; 83: 75-82.',
        'WHO Expert Committee on Leprosy: eighth report. WHO Tech Rep Ser 2012; 968: 1-61.',
        'Walker SL, et al. The role of thalidomide in ENL management. Lepr Rev 2007; 78: 197-215.',
        'Negera E, et al. Clinico-pathological features of ENL. PLoS Negl Trop Dis 2017; 11: e0006011.',
        'Kaufmann SHE, et al. Host-directed therapies for bacterial and viral infections. Nat Rev Drug Discov 2018; 17: 35-56.',
        'Hawn TR, et al. Countering antibiotic resistance with host-directed therapeutics. Immunol Rev 2015; 264: 344-362.',
        'Zumla A, Maeurer M. Host-directed therapies for MDR-TB. Clin Infect Dis 2015; 61: 1432-1438.',
        'Wallis RS, Hafner R. Advancing host-directed therapy for tuberculosis. Nat Rev Immunol 2015; 15: 255-263.',
        'Cole ST, et al. Massive gene decay in the leprosy bacillus. Nature 2001; 409: 1007-1011.',
        'de Mattos Barbosa MG, et al. IDO and iron required for M. leprae survival. Microbes Infect 2017; 19: 505-514.',
        'Palermo ML, et al. Increased Treg and down-regulatory molecules in LL. Am J Trop Med Hyg 2012; 86: 878-883.',
        'Truman RW, Krahenbuhl JL. Viable M. leprae as a research reagent. Int J Lepr 2001; 69: 1-12.',
        'Siddalingaiah HS. Multi-omics pipeline for TB HDT targets. Int J Tuberc Lung Dis 2024 (submitted).',
        'Wilkinson MD, et al. The FAIR guiding principles. Sci Data 2016; 3: 160018.',
        'Barrett T, et al. NCBI GEO: archive for functional genomics data sets. Nucleic Acids Res 2013; 41: D991-D995.',
        'Belone AF, et al. Genome-wide screening of mRNA expression in leprosy. Front Genet 2015; 6: 334.',
        'Montoya DJ, et al. Dual RNA-Seq of human leprosy lesions. Cell Rep 2019; 26: 3574-3585.',
        'Zavala K, et al. Blood transcriptomic profiles in leprosy. Sci Rep 2021; 11: 18715.',
        'Blischak JD, et al. Predicting susceptibility to tuberculosis. Sci Rep 2017; 7: 5702.',
        'Wu C, et al. BioGPS and MyGene.info. Nucleic Acids Res 2013; 41: D561-D565.',
        'Ochoa D, et al. Open Targets Platform. Nucleic Acids Res 2021; 49: D1302-D1310.',
        'Zdrazil B, et al. The ChEMBL Database in 2023. Nucleic Acids Res 2024; 52: D1180-D1192.',
        'Hunter JD. Matplotlib: a 2D graphics environment. Comput Sci Eng 2007; 9: 90-95.',
        'Thangaraju P, et al. Tofacitinib for chronic type II lepra reaction. Lepr Rev 2020; 91: 86-89.',
        'de Souza Sales J, et al. IDO role in LL immunosuppression. Clin Exp Immunol 2011; 165: 251-263.',
        'Bobosha K, et al. T-cell regulation in lepromatous leprosy. PLoS Negl Trop Dis 2014; 8: e2773.',
        'Cliff JM, et al. Human immune response to TB. Immunol Rev 2015; 264: 88-102.',
    ]
    
    for i, ref in enumerate(references):
        p = doc.add_paragraph()
        p.add_run(f'{i+1}. ').bold = True
        p.add_run(ref)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Save
    output_path = BASE_DIR / 'manuscripts' / 'Manuscript_Leprosy_HDT_COMPLETE.docx'
    doc.save(str(output_path))
    print(f'Created: {output_path}')
    print(f'Tables embedded: 3')
    print(f'Figures embedded: 4')

if __name__ == '__main__':
    create_manuscript()
