from docx import Document
from docx.shared import Pt

# Create Response to Reviewers DOCX (simplified)
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

doc.add_heading('Response to Reviewers', level=0)
doc.add_paragraph('Manuscript: An Integrated Multi-omics and Chemoinformatics Pipeline Identifies IDO1, PDL1, and JAK2 as Host-Directed Therapy Candidates for Leprosy')
doc.add_paragraph()

doc.add_heading('Summary of Revisions', level=1)
doc.add_paragraph('We thank both reviewers for their constructive reviews. We have addressed all 18 concerns and substantially strengthened the manuscript.')

doc.add_heading('Key Improvements:', level=2)
improvements = [
    'Expanded Methods with gene selection criteria, preprocessing, weight justification',
    'Systematic literature validation for ALL 50 targets (Supplementary Table S3)',
    'Sensitivity analysis: 10 weighting schemes tested (Supplementary Table S2)',
    'NEW SECTION: Immunological Context and HDT Timing',
    'NEW SECTION: Safety Considerations (autoimmune, TB reactivation, drug interactions)',
    'NEW SECTION: VEGFA Ranking vs Biological Plausibility',
    'NEW SECTION: Limitations',
    'Supplementary Table S1: All 50 genes with metadata'
]
for item in improvements:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

doc.add_heading('Response to Reviewer 1', level=1)

doc.add_heading('Major Concern 1: Gene Selection', level=2)
doc.add_paragraph('Reviewer: Insufficient detail on gene selection process')
doc.add_paragraph('Response: Added detailed criteria - genes must be DE in ≥2 datasets. Created Supp Table S1 with all metadata. Added Venn diagram reference.')

doc.add_heading('Major Concern 2: Scoring Weights', level=2)
doc.add_paragraph('Reviewer: Weights appear arbitrary')
doc.add_paragraph('Response: Added rationale for each weight. Performed sensitivity analysis (10 schemes). Top targets stable across variations.')

doc.add_heading('Major Concern 3: Compound Strategy', level=2)
doc.add_paragraph('Reviewer: TB overlap confusing')
doc.add_paragraph('Response: Clarified - 629 total compounds (217 leprosy-specific, 412 shared targets). Added Venn diagram.')

doc.add_heading('Major Concern 4: Literature Validation', level=2)
doc.add_paragraph('Reviewer: Only 3 targets validated')
doc.add_paragraph('Response: Systematic validation for ALL 50 targets. Created Supp Table S3. Key finding: BLK has ZERO publications - computational artifact.')

doc.add_heading('Major Concern 5: Methods Details', level=2)
doc.add_paragraph('Reviewer: Missing preprocessing, thresholds, API queries')
doc.add_paragraph('Response: Added preprocessing section (RMA, DESeq2, FDR). Justified thresholds. Specified API queries in Supp Methods.')

doc.add_page_break()

doc.add_heading('Response to Reviewer 2', level=1)

doc.add_heading('Major Concern 1: Immunological Context', level=2)
doc.add_paragraph('Reviewer: Reactions occur in different contexts')
doc.add_paragraph('Response: Added Section 4.1 - Type 1 vs Type 2 mechanisms. HDT timing: prevention/treatment/maintenance. Immune reconstitution risk discussed.')

doc.add_heading('Major Concern 2: Safety', level=2)
doc.add_paragraph('Reviewer: Significant safety concerns not discussed')
doc.add_paragraph('Response: Added Section 4.2 - Checkpoint inhibitors (autoimmune 10-20%), IDO1 (dual role), JAK inhibitors (TB reactivation, rifampicin interactions 50-70% reduction). Mitigation strategies.')

doc.add_heading('Major Concern 3: Target Biology', level=2)
doc.add_paragraph('Reviewer: Biology oversimplified')
doc.add_paragraph('Response: Expanded IDO1 dual role discussion. Specified PD-L1 cell types. Added Section 4.3 on VEGFA ranking (high druggability but limited validation).')

doc.add_page_break()

doc.add_heading('Conclusion', level=1)
doc.add_paragraph('All 18 concerns addressed. Manuscript strengthened with:')
items = [
    'Detailed methods and supplementary data',
    'Clinical and immunological context',
    'Comprehensive safety considerations',
    'Systematic validation for all targets',
    'Clear limitations'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Word count: 2,993 → 3,200 words')
doc.add_paragraph()
doc.add_paragraph('We hope the revised manuscript is suitable for publication.')
doc.add_paragraph()
doc.add_paragraph('Sincerely,')
doc.add_paragraph('Dr. Siddalingaiah H S')

doc.save('manuscripts/Response_to_Reviewers.docx')
print('Created Response_to_Reviewers.docx')
