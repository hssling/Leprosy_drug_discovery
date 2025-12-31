from docx import Document
from docx.shared import Pt

# Create Supplementary Materials DOCX
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

doc.add_heading('Supplementary Materials', level=0)
doc.add_paragraph('An Integrated Multi-omics and Chemoinformatics Pipeline Identifies IDO1, PDL1, and JAK2 as Host-Directed Therapy Candidates for Leprosy')
doc.add_paragraph()

# Table S1
doc.add_heading('Supplementary Table S1: Gene List', level=1)
doc.add_paragraph('All 50 genes with expression data, dataset coverage, and validation status.')
doc.add_paragraph('Available in CSV format: Supplementary_Table_S1.csv')

table = doc.add_table(rows=11, cols=5)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Gene'
header[1].text = 'Datasets'
header[2].text = 'PubMed'
header[3].text = 'Function'
header[4].text = 'Validation'

genes = [
    ['VEGFA', '4', '4', 'Angiogenesis', 'Limited'],
    ['PDL1', '4', '8', 'Immune checkpoint', 'Moderate'],
    ['IDO1', '4', '12', 'Tryptophan metabolism', 'Strong'],
    ['CD38', '3', '3', 'NADase enzyme', 'Limited'],
    ['IL10', '4', '18', 'Anti-inflammatory', 'Strong'],
    ['VDR', '3', '15', 'Vitamin D receptor', 'Strong'],
    ['BLK', '2', '0', 'B-cell kinase', 'None'],
    ['JAK2', '3', '3', 'JAK-STAT signaling', 'Limited'],
    ['CXCL10', '4', '6', 'Chemokine', 'Moderate'],
    ['CTLA4', '3', '2', 'Immune checkpoint', 'Limited']
]

for i, gene_data in enumerate(genes, 1):
    cells = table.rows[i].cells
    for j, value in enumerate(gene_data):
        cells[j].text = value

doc.add_page_break()

# Table S2
doc.add_heading('Supplementary Table S2: Sensitivity Analysis', level=1)
doc.add_paragraph('Testing 10 weighting schemes for composite scoring.')
doc.add_paragraph('Available in CSV format: Supplementary_Table_S2.csv')

table2 = doc.add_table(rows=6, cols=5)
table2.style = 'Light Grid Accent 1'

header2 = table2.rows[0].cells
header2[0].text = 'Scheme'
header2[1].text = 'Omics'
header2[2].text = 'Top 1'
header2[3].text = 'Top 2'
header2[4].text = 'Top 3'

schemes = [
    ['Original', '0.35', 'VEGFA', 'PDL1', 'IDO1'],
    ['Equal', '0.20', 'VEGFA', 'PDL1', 'IDO1'],
    ['Omics Heavy', '0.50', 'IDO1', 'CXCL10', 'PDL1'],
    ['Druggability', '0.20', 'VEGFA', 'PDL1', 'IL10'],
    ['Evidence', '0.25', 'PDL1', 'IDO1', 'VEGFA']
]

for i, scheme in enumerate(schemes, 1):
    cells = table2.rows[i].cells
    for j, value in enumerate(scheme):
        cells[j].text = value

doc.add_paragraph()
doc.add_paragraph('Key Finding: IDO1, PDL1, JAK2 consistently in top 15 across all schemes.')

doc.add_page_break()

# Table S3
doc.add_heading('Supplementary Table S3: Literature Validation', level=1)
doc.add_paragraph('Systematic PubMed search for all 50 targets.')
doc.add_paragraph('Search: Gene AND (leprosy OR M. leprae OR Hansen)')
doc.add_paragraph()

doc.add_heading('Categories:', level=2)
doc.add_paragraph('Strong: ≥10 publications', style='List Bullet')
doc.add_paragraph('Moderate: 5-9 publications', style='List Bullet')
doc.add_paragraph('Limited: 1-4 publications', style='List Bullet')
doc.add_paragraph('None: 0 publications', style='List Bullet')

doc.add_paragraph()
doc.add_heading('Summary:', level=2)
doc.add_paragraph('Strong: 8 targets (16%) - IDO1, IL10, TNF, VDR, IFNG, IL1B, IL6, TGFB1', style='List Bullet')
doc.add_paragraph('Moderate: 12 targets (24%)', style='List Bullet')
doc.add_paragraph('Limited: 24 targets (48%)', style='List Bullet')
doc.add_paragraph('None: 6 targets (12%) - BLK, IL23A, IL22, MARCO, MSR1, PIAS1, PIAS3', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Critical Finding: BLK (Rank 7) has ZERO publications - computational artifact.')

doc.save('manuscripts/Supplementary_Materials_REVISED.docx')
print('Created Supplementary_Materials_REVISED.docx')
