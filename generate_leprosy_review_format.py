"""
Generate Leprosy Review compliant manuscript DOCX
- Double-spaced
- Line numbers
- Left-aligned
- Page numbers
- Restructured abstract (no Background section)
- Separate title page
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from pathlib import Path
import re

BASE_DIR = Path(__file__).parent

def add_line_numbers(doc):
    """Add line numbers to document"""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = sectPr.find(qn('w:lnNumType'))
        if lnNumType is None:
            lnNumType = sectPr.makeelement(qn('w:lnNumType'))
            sectPr.append(lnNumType)
        lnNumType.set(qn('w:countBy'), '1')
        lnNumType.set(qn('w:restart'), 'continuous')

def set_double_spacing(paragraph):
    """Set paragraph to double spacing"""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

def add_formatted_run(para, text):
    """Add text with superscript handling"""
    parts = re.split(r'(\^\d+(?:,\d+)*\^)', text)
    for part in parts:
        if part.startswith('^') and part.endswith('^'):
            run = para.add_run(part[1:-1])
            run.font.superscript = True
        else:
            para.add_run(part)

# Generate Title Page
title_doc = Document()
title_doc.sections[0].page_height = Inches(11)
title_doc.sections[0].page_width = Inches(8.5)

# Title
title = title_doc.add_paragraph()
title_run = title.add_run('An Integrated Multi-omics and Chemoinformatics Pipeline Identifies IDO1, PDL1, and JAK2 as Host-Directed Therapy Candidates for Leprosy')
title_run.font.size = Pt(14)
title_run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_double_spacing(title)

title_doc.add_paragraph()

# Author
author = title_doc.add_paragraph()
author.add_run('Siddalingaiah H S').bold = True
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_double_spacing(author)

# Affiliation
aff = title_doc.add_paragraph()
aff.add_run('Department of Community Medicine')
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_double_spacing(aff)

aff2 = title_doc.add_paragraph()
aff2.add_run('Shridevi Institute of Medical Sciences and Research Hospital')
aff2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_double_spacing(aff2)

aff3 = title_doc.add_paragraph()
aff3.add_run('Tumkur – 572106, Karnataka, India')
aff3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_double_spacing(aff3)

title_doc.add_paragraph()

# Corresponding author
corr = title_doc.add_paragraph()
corr.add_run('Corresponding Author:').bold = True
set_double_spacing(corr)

corr2 = title_doc.add_paragraph()
corr2.add_run('Dr. Siddalingaiah H S')
set_double_spacing(corr2)

corr3 = title_doc.add_paragraph()
corr3.add_run('Email: hssling@yahoo.com')
set_double_spacing(corr3)

corr4 = title_doc.add_paragraph()
corr4.add_run('Phone: +91-8941087719')
set_double_spacing(corr4)

corr5 = title_doc.add_paragraph()
corr5.add_run('ORCID: 0000-0002-4771-8285')
set_double_spacing(corr5)

# Save title page
title_doc.save(str(BASE_DIR / 'manuscripts' / 'TitlePage_LeprosyReview.docx'))
print('Created TitlePage_LeprosyReview.docx')

# Note: Main manuscript formatting would require regenerating the entire manuscript
# with double-spacing and line numbers. This is a template showing the approach.

print('\\nNext steps:')
print('1. Regenerate main manuscript with double-spacing')
print('2. Add line numbers to main manuscript')
print('3. Ensure left alignment (not justified)')
print('4. Add page numbers')
print('5. Restructure abstract (remove Background section)')
